"""Tests for vector search (T4 part 3).

Covers:
  FR-ING-4 – Vector search with topK retrieval
  Matter-scoped vector search with ACL enforcement
  Anchor preservation in search results
"""

from pathlib import Path

import pytest
from docx import Document as DocxDocument

from app.config import Settings, get_settings
from app.main import app


@pytest.fixture
def sample_legal_docx(tmp_path: Path) -> Path:
    """Create a DOCX with legal-style content for testing."""
    docx_path = tmp_path / "legal_contract.docx"
    doc = DocxDocument()

    doc.add_paragraph(
        "ARTICLE 1: DEFINITIONS. For purposes of this agreement, the following definitions apply."
    )
    doc.add_paragraph(
        "ARTICLE 2: PAYMENT TERMS. The buyer shall pay the seller within thirty days of invoice."
    )
    doc.add_paragraph(
        "ARTICLE 3: CONFIDENTIALITY. All parties shall maintain strict confidentiality of proprietary information."
    )
    doc.add_paragraph(
        "ARTICLE 4: TERMINATION. Either party may terminate this agreement with written notice."
    )

    doc.save(docx_path)
    return docx_path


@pytest.fixture
def _override_settings_with_matter(tmp_path: Path):
    """Override settings for vector search tests."""
    storage_path = tmp_path / "storage"
    storage_path.mkdir(exist_ok=True)

    test_settings = Settings(
        user_matters={"alice": ["matter-1"], "bob": ["matter-2"]},
        storage_path=str(storage_path),
    )
    app.dependency_overrides[get_settings] = lambda: test_settings
    yield
    app.dependency_overrides.clear()


# Note: These tests require PostgreSQL with pgvector for full functionality.
# With SQLite, embeddings are stored but vector search operations are limited.
# Mark as skip if not using PostgreSQL.


@pytest.mark.skip(reason="Requires PostgreSQL with pgvector extension")
async def test_search_requires_authentication(client):
    resp = await client.post(
        "/v1/search",
        json={"matter_id": "matter-1", "query": "payment terms"},
    )
    assert resp.status_code == 401


@pytest.mark.skip(reason="Requires PostgreSQL with pgvector extension")
async def test_search_enforces_matter_acl(
    client, sample_legal_docx: Path, _override_settings_with_matter
):
    resp = await client.post(
        "/v1/search",
        json={"matter_id": "matter-99", "query": "payment"},
        headers={"X-User-Id": "alice"},
    )

    assert resp.status_code == 403
    assert "not authorized for this matter" in resp.json()["detail"]


@pytest.mark.skip(reason="Requires PostgreSQL with pgvector extension")
async def test_search_returns_relevant_chunks(
    client, sample_legal_docx: Path, _override_settings_with_matter
):
    with sample_legal_docx.open("rb") as f:
        upload_resp = await client.post(
            "/v1/documents/upload?matter_id=matter-1",
            files={"file": ("contract.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            headers={"X-User-Id": "alice"},
        )

    assert upload_resp.status_code == 200

    search_resp = await client.post(
        "/v1/search",
        json={"matter_id": "matter-1", "query": "payment terms invoice", "top_k": 3},
        headers={"X-User-Id": "alice"},
    )

    assert search_resp.status_code == 200
    data = search_resp.json()

    assert data["matter_id"] == "matter-1"
    assert data["query"] == "payment terms invoice"
    assert len(data["results"]) <= 3

    result = data["results"][0]
    assert "chunk_id" in result
    assert "doc_id" in result
    assert "chunk_text" in result
    assert "anchor_start" in result
    assert "anchor_end" in result
    assert "similarity_score" in result
    assert 0.0 <= result["similarity_score"] <= 1.0


@pytest.mark.skip(reason="Requires PostgreSQL with pgvector extension")
async def test_search_respects_top_k_parameter(
    client, sample_legal_docx: Path, _override_settings_with_matter
):
    with sample_legal_docx.open("rb") as f:
        await client.post(
            "/v1/documents/upload?matter_id=matter-1",
            files={"file": ("contract.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            headers={"X-User-Id": "alice"},
        )

    resp_2 = await client.post(
        "/v1/search",
        json={"matter_id": "matter-1", "query": "agreement", "top_k": 2},
        headers={"X-User-Id": "alice"},
    )

    resp_5 = await client.post(
        "/v1/search",
        json={"matter_id": "matter-1", "query": "agreement", "top_k": 5},
        headers={"X-User-Id": "alice"},
    )

    assert len(resp_2.json()["results"]) <= 2
    assert len(resp_5.json()["results"]) <= 5


@pytest.mark.skip(reason="Requires PostgreSQL with pgvector extension")
async def test_search_isolates_matters(
    client, sample_legal_docx: Path, _override_settings_with_matter
):
    with sample_legal_docx.open("rb") as f1:
        await client.post(
            "/v1/documents/upload?matter_id=matter-1",
            files={"file": ("contract1.docx", f1, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            headers={"X-User-Id": "alice"},
        )

    doc2_path = sample_legal_docx.parent / "other_doc.docx"
    doc = DocxDocument()
    doc.add_paragraph("Completely unrelated content about gardening and plants.")
    doc.save(doc2_path)

    with doc2_path.open("rb") as f2:
        await client.post(
            "/v1/documents/upload?matter_id=matter-2",
            files={"file": ("other.docx", f2, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            headers={"X-User-Id": "bob"},
        )

    alice_resp = await client.post(
        "/v1/search",
        json={"matter_id": "matter-1", "query": "payment terms"},
        headers={"X-User-Id": "alice"},
    )

    bob_resp = await client.post(
        "/v1/search",
        json={"matter_id": "matter-2", "query": "gardening"},
        headers={"X-User-Id": "bob"},
    )

    assert alice_resp.status_code == 200
    assert bob_resp.status_code == 200

    alice_results = alice_resp.json()["results"]
    bob_results = bob_resp.json()["results"]

    if alice_results:
        assert "payment" in alice_results[0]["chunk_text"].lower() or "terms" in alice_results[0]["chunk_text"].lower()

    if bob_results:
        assert "garden" in bob_results[0]["chunk_text"].lower() or "plant" in bob_results[0]["chunk_text"].lower()


@pytest.mark.skip(reason="Requires PostgreSQL with pgvector extension")
async def test_search_results_include_anchors(
    client, sample_legal_docx: Path, _override_settings_with_matter
):
    with sample_legal_docx.open("rb") as f:
        await client.post(
            "/v1/documents/upload?matter_id=matter-1",
            files={"file": ("contract.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            headers={"X-User-Id": "alice"},
        )

    resp = await client.post(
        "/v1/search",
        json={"matter_id": "matter-1", "query": "definitions"},
        headers={"X-User-Id": "alice"},
    )

    results = resp.json()["results"]
    assert len(results) > 0

    result = results[0]
    assert result["anchor_start"]
    assert result["anchor_end"]
    assert result["anchor_start"].isdigit()
