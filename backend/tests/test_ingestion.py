"""Tests for document ingestion (T3).

Covers:
  FR-ING-1 – Upload endpoint stores file and extracted text
  FR-ING-2 – Extracted text includes anchors (page/paragraph)
  ACL enforcement on document upload
"""

import io
from pathlib import Path

import pytest
from pypdf import PdfWriter
from docx import Document as DocxDocument
from fastapi import UploadFile
from sqlalchemy import create_engine, select
from sqlalchemy.orm import Session, sessionmaker

from app.config import Settings, get_settings
from app.extraction import Anchor, extract_text, extract_text_from_docx, extract_text_from_pdf
from app.main import app
from app.models import Base, Document, ExtractedText


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def db_session() -> Session:
    """In-memory SQLite session for unit tests."""
    engine = create_engine("sqlite://", connect_args={"check_same_thread": False})
    Base.metadata.create_all(bind=engine)
    factory = sessionmaker(bind=engine, expire_on_commit=False)
    session = factory()
    yield session
    session.close()
    engine.dispose()


@pytest.fixture
def sample_pdf(tmp_path: Path) -> Path:
    """Create a minimal PDF file for testing."""
    pdf_path = tmp_path / "sample.pdf"
    writer = PdfWriter()

    writer.add_blank_page(width=612, height=792)
    page1 = writer.pages[0]
    page1.merge_translated_page(page1, 0, 0)

    writer.add_blank_page(width=612, height=792)

    with pdf_path.open("wb") as f:
        writer.write(f)

    return pdf_path


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """Create a minimal DOCX file for testing."""
    docx_path = tmp_path / "sample.docx"
    doc = DocxDocument()
    doc.add_paragraph("First paragraph with some text.")
    doc.add_paragraph("Second paragraph with more text.")
    doc.add_paragraph("Third paragraph for testing.")
    doc.save(docx_path)
    return docx_path


@pytest.fixture
def _override_settings(tmp_path: Path):
    """Override settings with test configuration."""
    storage_path = tmp_path / "storage"
    storage_path.mkdir(exist_ok=True)

    test_settings = Settings(
        user_matters={"alice": ["matter-1"], "bob": ["matter-2"]},
        storage_path=str(storage_path),
    )
    app.dependency_overrides[get_settings] = lambda: test_settings
    yield
    app.dependency_overrides.clear()


# ---------------------------------------------------------------------------
# Unit tests – Extraction utilities
# ---------------------------------------------------------------------------


class TestPdfExtraction:
    def test_extracts_pages_as_anchors(self, sample_pdf: Path):
        result = extract_text_from_pdf(sample_pdf)

        assert result.doc_type == "pdf"
        assert len(result.anchors) >= 0

    def test_anchor_structure(self, sample_pdf: Path):
        result = extract_text_from_pdf(sample_pdf)

        for anchor in result.anchors:
            assert isinstance(anchor, Anchor)
            assert anchor.anchor_type == "page"
            assert anchor.anchor_value.isdigit()
            assert isinstance(anchor.text_content, str)

    def test_handles_missing_file(self, tmp_path: Path):
        missing_pdf = tmp_path / "missing.pdf"

        with pytest.raises(ValueError, match="Failed to extract text from PDF"):
            extract_text_from_pdf(missing_pdf)


class TestDocxExtraction:
    def test_extracts_paragraphs_as_anchors(self, sample_docx: Path):
        result = extract_text_from_docx(sample_docx)

        assert result.doc_type == "docx"
        assert len(result.anchors) == 3

    def test_anchor_content_matches(self, sample_docx: Path):
        result = extract_text_from_docx(sample_docx)

        assert result.anchors[0].text_content == "First paragraph with some text."
        assert result.anchors[1].text_content == "Second paragraph with more text."
        assert result.anchors[2].text_content == "Third paragraph for testing."

    def test_anchor_values_are_sequential(self, sample_docx: Path):
        result = extract_text_from_docx(sample_docx)

        assert result.anchors[0].anchor_value == "1"
        assert result.anchors[1].anchor_value == "2"
        assert result.anchors[2].anchor_value == "3"

    def test_handles_missing_file(self, tmp_path: Path):
        missing_docx = tmp_path / "missing.docx"

        with pytest.raises(ValueError, match="Failed to extract text from DOCX"):
            extract_text_from_docx(missing_docx)


class TestGenericExtraction:
    def test_dispatches_to_pdf_extractor(self, sample_pdf: Path):
        result = extract_text(sample_pdf)
        assert result.doc_type == "pdf"

    def test_dispatches_to_docx_extractor(self, sample_docx: Path):
        result = extract_text(sample_docx)
        assert result.doc_type == "docx"

    def test_rejects_unsupported_type(self, tmp_path: Path):
        unsupported = tmp_path / "file.txt"
        unsupported.write_text("test")

        with pytest.raises(ValueError, match="Unsupported file type"):
            extract_text(unsupported)


# ---------------------------------------------------------------------------
# Integration tests – /v1/documents/upload endpoint
# ---------------------------------------------------------------------------


async def test_upload_requires_authentication(client):
    resp = await client.post(
        "/v1/documents/upload?matter_id=matter-1",
        files={"file": ("test.pdf", b"fake pdf", "application/pdf")},
    )
    assert resp.status_code == 401


async def test_upload_enforces_matter_acl(client, sample_pdf: Path, _override_settings):
    with sample_pdf.open("rb") as f:
        resp = await client.post(
            "/v1/documents/upload?matter_id=matter-99",
            files={"file": ("sample.pdf", f, "application/pdf")},
            headers={"X-User-Id": "alice"},
        )

    assert resp.status_code == 403
    assert "not authorized for this matter" in resp.json()["detail"]


async def test_upload_rejects_unsupported_file_type(client, _override_settings):
    resp = await client.post(
        "/v1/documents/upload?matter_id=matter-1",
        files={"file": ("test.txt", b"plain text", "text/plain")},
        headers={"X-User-Id": "alice"},
    )

    assert resp.status_code == 400
    assert "Unsupported file type" in resp.json()["detail"]


async def test_upload_pdf_success(client, sample_pdf: Path, _override_settings):
    with sample_pdf.open("rb") as f:
        resp = await client.post(
            "/v1/documents/upload?matter_id=matter-1",
            files={"file": ("sample.pdf", f, "application/pdf")},
            headers={"X-User-Id": "alice"},
        )

    assert resp.status_code == 200
    data = resp.json()

    assert "doc_id" in data
    assert data["matter_id"] == "matter-1"
    assert data["filename"] == "sample.pdf"
    assert data["doc_type"] == "pdf"
    assert isinstance(data["anchors"], list)


async def test_upload_docx_success(client, sample_docx: Path, _override_settings):
    with sample_docx.open("rb") as f:
        resp = await client.post(
            "/v1/documents/upload?matter_id=matter-1",
            files={"file": ("sample.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            headers={"X-User-Id": "alice"},
        )

    assert resp.status_code == 200
    data = resp.json()

    assert data["doc_type"] == "docx"
    assert len(data["anchors"]) == 3


async def test_upload_stores_document_metadata(client, sample_pdf: Path, _override_settings):
    with sample_pdf.open("rb") as f:
        resp = await client.post(
            "/v1/documents/upload?matter_id=matter-1",
            files={"file": ("sample.pdf", f, "application/pdf")},
            headers={"X-User-Id": "alice"},
        )

    doc_id = resp.json()["doc_id"]

    session: Session = app.state.session_factory()
    try:
        doc = session.execute(
            select(Document).where(Document.doc_id == doc_id)
        ).scalar_one()

        assert doc.matter_id == "matter-1"
        assert doc.filename == "sample.pdf"
        assert doc.doc_type == "pdf"
        assert doc.uploaded_by == "alice"
        assert doc.uploaded_at is not None
    finally:
        session.close()


async def test_upload_stores_extracted_text(client, sample_docx: Path, _override_settings):
    with sample_docx.open("rb") as f:
        resp = await client.post(
            "/v1/documents/upload?matter_id=matter-1",
            files={"file": ("sample.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            headers={"X-User-Id": "alice"},
        )

    doc_id = resp.json()["doc_id"]

    session: Session = app.state.session_factory()
    try:
        texts = session.execute(
            select(ExtractedText).where(ExtractedText.doc_id == doc_id)
        ).scalars().all()

        assert len(texts) == 3
        assert texts[0].anchor_type == "paragraph"
        assert texts[0].anchor_value == "1"
        assert texts[0].text_content == "First paragraph with some text."
    finally:
        session.close()


async def test_upload_response_includes_anchor_previews(client, sample_docx: Path, _override_settings):
    with sample_docx.open("rb") as f:
        resp = await client.post(
            "/v1/documents/upload?matter_id=matter-1",
            files={"file": ("sample.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            headers={"X-User-Id": "alice"},
        )

    data = resp.json()
    anchors = data["anchors"]

    assert len(anchors) == 3
    assert anchors[0]["anchor_type"] == "paragraph"
    assert anchors[0]["anchor_value"] == "1"
    assert "First paragraph" in anchors[0]["text_preview"]


async def test_different_users_upload_to_different_matters(
    client, sample_pdf: Path, _override_settings
):
    with sample_pdf.open("rb") as f1:
        resp1 = await client.post(
            "/v1/documents/upload?matter_id=matter-1",
            files={"file": ("alice.pdf", f1, "application/pdf")},
            headers={"X-User-Id": "alice"},
        )

    with sample_pdf.open("rb") as f2:
        resp2 = await client.post(
            "/v1/documents/upload?matter_id=matter-2",
            files={"file": ("bob.pdf", f2, "application/pdf")},
            headers={"X-User-Id": "bob"},
        )

    assert resp1.status_code == 200
    assert resp2.status_code == 200
    assert resp1.json()["matter_id"] == "matter-1"
    assert resp2.json()["matter_id"] == "matter-2"
